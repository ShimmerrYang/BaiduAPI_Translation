# -*- coding: utf-8 -*-

import requests
import random
from hashlib import md5
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# name = input('请输入Word的文件名：')
qianzhui = 'demo'
name = qianzhui+'.docx'

# User: YzH
appid = '20210325000742262'
appkey = 'nBZkYdALLeoNfFhFxcfT'

# For list of language codes, please refer to `https://api.fanyi.baidu.com/doc/21`
from_lang = 'en'
to_lang = 'zh'

# -------------------------------参数设置完成-------------------------------

endpoint = 'http://api.fanyi.baidu.com'
path = '/api/trans/vip/translate'
url = endpoint + path


# Generate salt and sign
def make_md5(s, encoding='utf-8'):
    return md5(s.encode(encoding)).hexdigest()


doc = Document(name)
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(12)

# 开始循环提交翻译
step = 0
for query in doc.paragraphs:

    query = query.text  # 把类转换成要翻译的字符串
    if query == '':
        step = step + 1
        continue
    if query == 'References' or query == 'Reference' or query == 'references' or query == 'reference':
        break

    print('第{}段读取完成，开始提交数据...'.format(step+1))
    print(query)

    salt = random.randint(32768, 65536)
    sign = make_md5(appid + query + str(salt) + appkey)

    # Build request
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    payload = {'appid': appid, 'q': query, 'from': from_lang, 'to': to_lang, 'salt': salt, 'sign': sign}

    # Send request
    r = requests.post(url, params=payload, headers=headers)
    result = r.json()

    print(result)  # 对result可视化显示

    # Save in a new word
    doc.paragraphs[step].text = query + result['trans_result'][0]['dst']

    doc.save(qianzhui+'_翻译版.docx')
    step = step + 1
    time.sleep(1.3)

print('翻译完成！！！')

# 现存bug，换行符超时，超链接丢失